import os
import uuid
import json
import mimetypes
import base64
import requests
import re
from urllib.parse import quote, urljoin
from PIL import Image
import io
try:
    import PyPDF2
    print("✅ PyPDF2 importado com sucesso")
except ImportError as e:
    print(f"❌ Erro ao importar PyPDF2: {e}")
    PyPDF2 = None

try:
    import docx
    print("✅ python-docx importado com sucesso")
except ImportError as e:
    print(f"❌ Erro ao importar python-docx: {e}")
    docx = None

try:
    from bs4 import BeautifulSoup
    print("✅ BeautifulSoup importado com sucesso")
except ImportError as e:
    print(f"❌ Erro ao importar BeautifulSoup: {e}")
    BeautifulSoup = None
from flask import Flask, request, jsonify, send_from_directory, session
from flask_cors import CORS
try:
    from huggingface_hub import InferenceClient
    print("✅ huggingface_hub importado com sucesso")
except ImportError as e:
    print(f"❌ Erro ao importar huggingface_hub: {e}")
    InferenceClient = None
from dotenv import load_dotenv # Importe esta linha para carregar variáveis de ambiente

# --- CARREGA VARIÁVEIS DE AMBIENTE ---
# Isso deve estar no topo do seu arquivo app.py, antes de usar os.getenv para as chaves
load_dotenv()

# --- VERIFICAÇÕES DO SISTEMA ---
print("=== INICIANDO A.E.M.I ===")
print(f"Python: {__import__('sys').version}")
print(f"Flask: {__import__('flask').__version__}")
print(f"Diretório atual: {os.getcwd()}")
print(f"Arquivos no diretório: {os.listdir('.')}")
print("=========================") 

# --- INICIALIZAÇÃO DO FLASK ---
app = Flask(__name__)

# Permite apenas o domínio do GitHub Pages do seu projeto e o endereço do backend Render
CORS(app, supports_credentials=True, origins=[
    "https://jonathan0078.github.io",
    "https://aemi.onrender.com"
])

# Carrega as chaves da aplicação a partir de variáveis de ambiente
FLASK_SECRET_KEY = os.getenv("FLASK_SECRET_KEY", "dev-key-change-in-production")
HUGGING_FACE_TOKEN = os.getenv("HF_TOKEN")
# VARIÁVEIS PARA GOOGLE CUSTOM SEARCH API
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_CSE_ID = os.getenv("GOOGLE_CSE_ID")

# Validação das chaves
if not HUGGING_FACE_TOKEN:
    print("AVISO: Token da Hugging Face não configurado. Chat com IA não funcionará.")
    print("Configure a variável de ambiente HF_TOKEN no Replit Secrets.")

if not GOOGLE_API_KEY:
    print("AVISO: GOOGLE_API_KEY não configurada. Funcionalidade de busca web não funcionará.")
    print("Configure as variáveis GOOGLE_API_KEY e GOOGLE_CSE_ID no Replit Secrets.")
    print("Siga: https://developers.google.com/custom-search/v1/introduction")

if not GOOGLE_CSE_ID:
    print("AVISO: GOOGLE_CSE_ID não configurado. Configure seu próprio CSE ID.")
    print("Crie um CSE em: https://cse.google.com/")

app.secret_key = FLASK_SECRET_KEY

# --- CONFIGURAÇÕES DA BASE DE CONHECIMENTO ---
KB_DIR = os.path.join(os.path.dirname(__file__), 'knowledge_base')
KB_DB = os.path.join(KB_DIR, 'kb_data.json')
os.makedirs(KB_DIR, exist_ok=True)

def load_kb():
    if not os.path.exists(KB_DB):
        return []
    try:
        with open(KB_DB, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Erro ao carregar KB: {e}")
        return []

def save_kb(data):
    try:
        with open(KB_DB, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Erro ao salvar KB: {e}")

# --- FUNÇÃO DE BUSCA NA INTERNET ---
def google_search_api(query, num_results=3):
    """
    Faz uma busca na web usando a Google Custom Search JSON API.
    """
    if not GOOGLE_API_KEY or not GOOGLE_CSE_ID:
        print("Erro: Chave de API do Google ou CSE ID não configurados.")
        return {"error": "API do Google não configurada. Configure GOOGLE_API_KEY e GOOGLE_CSE_ID no Replit Secrets."}

    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": GOOGLE_API_KEY,
        "cx": GOOGLE_CSE_ID,
        "q": query,
        "num": min(num_results, 10),  # Google limita a 10 resultados por requisição
        "safe": "active",  # Filtro de conteúdo seguro
        "lr": "lang_pt",  # Preferência por resultados em português
        "hl": "pt"  # Interface em português
    }

    try:
        print(f"Fazendo requisição para Google API com query: {query}")
        response = requests.get(url, params=params, timeout=15)
        
        # Log detalhado do status
        print(f"Status da resposta: {response.status_code}")
        
        if response.status_code == 403:
            error_data = response.json() if response.headers.get('content-type', '').startswith('application/json') else {}
            error_message = error_data.get('error', {}).get('message', 'Acesso negado')
            print(f"Erro 403: {error_message}")
            return {"error": f"API do Google: {error_message}. Verifique sua chave API e configurações."}
        
        if response.status_code == 429:
            return {"error": "Limite de requisições da API do Google excedido. Tente novamente mais tarde."}
        
        response.raise_for_status()
        search_results = response.json()
        
        print(f"Resposta da API recebida. Keys: {search_results.keys()}")

        results_list = []
        if "items" in search_results:
            print(f"Encontrados {len(search_results['items'])} resultados")
            for item in search_results["items"]:
                results_list.append({
                    "title": item.get("title", "Sem título"),
                    "link": item.get("link", ""),
                    "snippet": item.get("snippet", "Sem descrição")
                })
        else:
            print("Nenhum item encontrado na resposta da API")
        
        return {"results": results_list, "query": query}
        
    except requests.exceptions.Timeout:
        print("Timeout na requisição à API do Google")
        return {"error": "Timeout na pesquisa. Tente novamente."}
    except requests.exceptions.RequestException as e:
        print(f"Erro na requisição à API de busca do Google: {e}")
        return {"error": f"Erro de conexão com a API do Google: {str(e)}"}
    except ValueError as e:
        print(f"Erro ao processar JSON da resposta: {e}")
        return {"error": "Erro ao processar resposta da API do Google"}
    except Exception as e:
        print(f"Erro inesperado ao processar busca do Google: {e}")
        return {"error": f"Erro inesperado: {str(e)}"}

# --- ROTA PARA TESTAR A BUSCA DO GOOGLE ---
@app.route('/api/search', methods=['GET'])
def perform_search():
    query = request.args.get('q', '')
    if not query:
        return jsonify({"error": "Parâmetro 'q' (query) é obrigatório."}), 400
    
    num_results = int(request.args.get('num', 5))
    results = google_search_api(query, num_results)
    return jsonify(results)

# --- ROTAS DA BASE DE CONHECIMENTO (RESTO DO SEU CÓDIGO) ---
@app.route('/kb/upload', methods=['POST'])
def kb_upload():
    # ... o resto do seu código para kb_upload ...
    # (Copie e cole a partir daqui o que você já tinha no seu app.py)
    try:
        kb = load_kb()
        # ... seu código continua aqui ...

        # Exemplo de como você usaria a busca dentro de uma função da sua IA:
        # Quando o usuário faz uma pergunta que a IA não sabe responder diretamente,
        # ou que exige informação atualizada, você chamaria:
        # search_query = "informação sobre " + pergunta_do_usuario
        # search_results = google_search_api(search_query)
        # return jsonify({"response": "Encontrei isto: " + str(search_results)}) # Adaptar a resposta

        
        # Upload de arquivo
        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']
            ext = os.path.splitext(file.filename)[1].lower()
            file_id = str(uuid.uuid4())
            filename = f"{file_id}{ext}"
            file_path = os.path.join(KB_DIR, filename)
            file.save(file_path)
            
            # Extração de texto
            extracted_text = ''
            if ext == '.pdf' and PyPDF2:
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            extracted_text += page.extract_text() or ''
                except Exception as e:
                    print(f"Erro ao extrair PDF: {e}")
                    extracted_text = ''
            elif ext in ['.docx'] and docx:
                try:
                    doc = docx.Document(file_path)
                    extracted_text = '\n'.join([p.text for p in doc.paragraphs])
                except Exception as e:
                    print(f"Erro ao extrair DOCX: {e}")
                    extracted_text = ''
            elif ext in ['.txt', '.md', '.csv']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        extracted_text = f.read()
                except Exception as e:
                    print(f"Erro ao extrair texto: {e}")
                    extracted_text = ''
            
            kb.append({
                'id': file_id,
                'type': 'file',
                'name': file.filename,
                'filename': filename,
                'content': extracted_text[:20000] if extracted_text else ''
            })
        
        # Cadastro de FAQ/manual (texto)
        faq = request.form.get('faq', '').strip()
        if faq:
            faq_id = str(uuid.uuid4())
            kb.append({
                'id': faq_id,
                'type': 'faq',
                'name': faq[:40] + ('...' if len(faq) > 40 else ''),
                'content': faq
            })
        
        save_kb(kb)
        return jsonify({'success': True})
        
    except Exception as e:
        print(f"Erro ao fazer upload para KB: {e}")
        return jsonify({'error': 'Erro ao processar upload'}), 500

@app.route('/kb/list')
def kb_list():
    try:
        kb = load_kb()
        q = request.args.get('q', '').strip().lower()
        if q:
            kb = [item for item in kb if q in item.get('name','').lower() or q in item.get('content','').lower()]
        return jsonify(kb)
    except Exception as e:
        print(f"Erro ao listar KB: {e}")
        return jsonify([])

@app.route('/kb/remove/<item_id>', methods=['DELETE'])
def kb_remove(item_id):
    try:
        kb = load_kb()
        item = next((i for i in kb if i['id'] == item_id), None)
        if not item:
            return jsonify({'error': 'Item não encontrado'}), 404
        
        kb = [i for i in kb if i['id'] != item_id]
        
        if item['type'] == 'file' and 'filename' in item:
            try:
                os.remove(os.path.join(KB_DIR, item['filename']))
            except Exception:
                pass
        
        save_kb(kb)
        return jsonify({'success': True})
        
    except Exception as e:
        print(f"Erro ao remover item da KB: {e}")
        return jsonify({'error': 'Erro ao remover item'}), 500

@app.route('/kb/download/<item_id>')
def kb_download(item_id):
    try:
        kb = load_kb()
        item = next((i for i in kb if i['id'] == item_id and i['type'] == 'file'), None)
        if not item:
            return 'Arquivo não encontrado', 404
        return send_from_directory(KB_DIR, item['filename'], as_attachment=True, download_name=item['name'])
    except Exception as e:
        print(f"Erro ao fazer download: {e}")
        return 'Erro interno', 500

# --- CONSTANTES E CONFIGURAÇÕES DO CHAT ---
SYSTEM_PROMPT = "Você é a A.E.M.I, uma IA especialista em manutenção industrial e um projeto do canal 'Manutenção Industrial ARQUIVOS'. Seja direta e objetiva. Responda apenas a perguntas relacionadas a este domínio. Se a pergunta não for sobre manutenção industrial, diga que você só pode ajudar com tópicos relacionados à manutenção industrial."
MAX_HISTORY_LENGTH = 10

# --- FUNÇÕES DE PESQUISA NA INTERNET ---
def search_internet(query, max_results=5):
    """Pesquisa na internet usando Google Custom Search API."""
    print(f"Iniciando pesquisa para: {query}")
    
    # Primeiro tenta usar Google API
    google_results = google_search_api(query, max_results)
    
    if "error" not in google_results and "results" in google_results:
        print(f"Google API retornou {len(google_results['results'])} resultados")
        return google_results
    
    print(f"Google API falhou: {google_results.get('error', 'Erro desconhecido')}")
    
    # Fallback simples se Google API falhar
    return {
        "error": "Serviço de pesquisa temporariamente indisponível. Tente reformular sua pergunta ou me forneça mais contexto.",
        "query": query
    }

def extract_page_content(url, max_chars=2000):
    """Extrai conteúdo de uma página web para análise inteligente."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none'
        }
        
        response = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        
        if response.status_code == 200 and BeautifulSoup:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove elementos que não agregam conteúdo
            for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside', 
                               'iframe', 'noscript', 'form', 'button', 'input', 'select',
                               'meta', 'link', 'br', 'hr']):
                element.decompose()
            
            # Remove divs de publicidade e navegação
            for element in soup.find_all(['div', 'section'], {'class': re.compile(
                r'(ad|advertisement|sidebar|menu|nav|footer|header|social|share|comment|related)', re.I)}):
                element.decompose()
            
            # Busca conteúdo principal em ordem de prioridade
            content_selectors = [
                'article',
                '[role="main"]',
                'main',
                '.content',
                '.post-content', 
                '.entry-content',
                '.article-content',
                '.text-content',
                '#content',
                '#main-content',
                '.main-content'
            ]
            
            main_content = None
            for selector in content_selectors:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            # Se não encontrou área específica, usa o body
            if not main_content:
                main_content = soup.find('body')
            
            if not main_content:
                main_content = soup
            
            # Extrai texto de forma inteligente
            text_parts = []
            
            # Prioriza parágrafos, títulos e listas
            for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'blockquote']):
                text = element.get_text().strip()
                if text and len(text) > 10:  # Ignora textos muito curtos
                    text_parts.append(text)
            
            # Se não encontrou elementos estruturados, pega texto geral
            if not text_parts:
                text = main_content.get_text()
                lines = (line.strip() for line in text.splitlines())
                text_parts = [line for line in lines if line and len(line) > 10]
            
            # Remove duplicatas mantendo ordem
            seen = set()
            unique_parts = []
            for part in text_parts:
                if part not in seen:
                    seen.add(part)
                    unique_parts.append(part)
            
            # Junta o texto final
            full_text = ' '.join(unique_parts)
            
            # Limpa caracteres especiais e espaços excessivos
            full_text = re.sub(r'\s+', ' ', full_text)  # Múltiplos espaços → 1 espaço
            full_text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\%\$\€\£\°\+\=\/\\\[\]]', '', full_text)  # Remove caracteres especiais
            
            # Garante que termina numa frase completa se possível
            if len(full_text) > max_chars:
                truncated = full_text[:max_chars]
                last_period = truncated.rfind('.')
                last_exclamation = truncated.rfind('!')
                last_question = truncated.rfind('?')
                
                # Pega o último ponto, exclamação ou interrogação
                last_sentence_end = max(last_period, last_exclamation, last_question)
                
                if last_sentence_end > max_chars * 0.7:  # Se estiver pelo menos em 70% do texto
                    full_text = truncated[:last_sentence_end + 1]
                else:
                    full_text = truncated
            
            return full_text.strip()
        
        return ""
        
    except requests.exceptions.Timeout:
        print(f"Timeout ao acessar {url}")
        return ""
    except requests.exceptions.RequestException as e:
        print(f"Erro de conexão ao acessar {url}: {e}")
        return ""
    except Exception as e:
        print(f"Erro inesperado ao extrair conteúdo de {url}: {e}")
        return ""

def should_search_internet(message):
    """Determina se a mensagem requer pesquisa na internet com IA mais inteligente."""
    search_triggers = [
        # Gatilhos diretos de pesquisa
        'pesquisar', 'buscar', 'procurar', 'pesquise', 'busque', 'procure',
        'google', 'internet', 'web', 'online',
        
        # Informações atuais/temporais
        'últimas', 'recente', 'atual', 'hoje', 'agora', 'notícias', 'novo', 'nova',
        'que dia é', 'data atual', 'horário', 'ano', '2024', '2025',
        
        # Informações comerciais
        'preço', 'valor', 'custo', 'onde comprar', 'fornecedor', 'venda', 'vender',
        'loja', 'mercado', 'empresa', 'fabricante', 'marca', 'modelo',
        
        # Especificações técnicas atuais
        'especificação', 'datasheet', 'manual', 'catálogo', 'norma', 'nbr', 'iso', 'abnt',
        'regulamento', 'lei', 'certificação',
        
        # Localização e contato
        'endereço', 'telefone', 'contato', 'site', 'website', 'email',
        'onde fica', 'localização',
        
        # Comparações e análises
        'comparar', 'diferença', 'melhor', 'pior', 'vantagem', 'desvantagem',
        'review', 'avaliação', 'opinião',
        
        # Tendências e novidades
        'tendência', 'inovação', 'tecnologia', 'lançamento', 'novidade',
        'estatística', 'dados', 'relatório'
    ]
    
    message_lower = message.lower()
    
    # 1. Gatilhos diretos sempre fazem pesquisa
    if any(trigger in message_lower for trigger in search_triggers):
        print(f"✅ Gatilho de pesquisa encontrado: '{message}'")
        return True
    
    # 2. Padrões que sempre requerem informação atual
    always_search_patterns = [
        'que dia é', 'qual a data', 'data de hoje', 'que horas',
        'quando foi', 'em que ano', 'quantos anos',
        'onde comprar', 'qual empresa', 'qual fabricante', 'quem fabrica',
        'onde encontrar', 'qual o site', 'como contactar'
    ]
    
    if any(pattern in message_lower for pattern in always_search_patterns):
        print(f"✅ Padrão de pesquisa obrigatório: '{message}'")
        return True
    
    # 3. Perguntas sobre informações específicas que podem precisar de dados atuais
    question_words = ['qual', 'quais', 'como', 'onde', 'quando', 'quanto', 'quantos', 
                      'por que', 'porque', 'existe', 'tem', 'há', 'possui']
    
    if any(qw in message_lower for qw in question_words):
        # Palavras que indicam necessidade de informação específica/atual
        specific_info_words = [
            # Produtos e equipamentos específicos
            'modelo', 'versão', 'especificação', 'características',
            
            # Informações comerciais
            'preço', 'custo', 'valor', 'disponível', 'estoque',
            
            # Localização e fornecedores
            'empresa', 'fabricante', 'fornecedor', 'distribuidor',
            'loja', 'vendedor', 'representante',
            
            # Informações técnicas atuais
            'norma', 'regulamento', 'certificação', 'aprovação',
            'compatível', 'recomendado', 'aprovado',
            
            # Comparações
            'melhor', 'diferença', 'vantagem', 'comparação',
            
            # Informações temporais
            'novo', 'recente', 'atual', 'último', 'atualizado'
        ]
        
        if any(word in message_lower for word in specific_info_words):
            print(f"✅ Pergunta específica que pode precisar de dados atuais: '{message}'")
            return True
    
    # 4. Mensagens sobre equipamentos que podem precisar de info específica
    equipment_words = ['motor', 'bomba', 'válvula', 'sensor', 'rolamento', 'bearing',
                      'correia', 'belt', 'engrenagem', 'gear', 'compressor', 'turbina']
    
    if any(eq in message_lower for eq in equipment_words):
        # Se menciona equipamento E pede informação específica
        specific_requests = ['especificação', 'manual', 'datasheet', 'fabricante',
                           'onde comprar', 'preço', 'modelo', 'versão']
        
        if any(req in message_lower for req in specific_requests):
            print(f"✅ Pergunta sobre equipamento específico: '{message}'")
            return True
    
    # 5. Se a mensagem é muito longa e parece ser uma pergunta complexa
    if len(message) > 50 and '?' in message:
        complex_indicators = ['detalhes', 'informações', 'dados', 'explicação',
                            'procedimento', 'processo', 'método', 'técnica']
        
        if any(ind in message_lower for ind in complex_indicators):
            print(f"✅ Pergunta complexa que pode se beneficiar de pesquisa: '{message}'")
            return True
    
    print(f"❌ Não requer pesquisa: '{message}'")
    return False

def analyze_search_content(search_data, original_query):
    """Analisa o conteúdo dos resultados de pesquisa e gera uma resposta elaborada como ChatGPT/Gemini."""
    if "error" in search_data:
        return f"🔍 **Pesquisa na Internet**\n\n❌ {search_data['error']}\n\nComo alternativa, posso ajudar com base no meu conhecimento sobre manutenção industrial."
    
    results = search_data.get("results", [])
    if not results:
        return f"🔍 **Pesquisa na Internet**\n\n🚫 Nenhum resultado encontrado para: \"{original_query}\"\n\nComo alternativa, posso ajudar com base no meu conhecimento sobre manutenção industrial."
    
    print(f"🔍 Analisando {len(results)} resultados para: {original_query}")
    
    # EXTRAI CONTEÚDO REAL DAS PÁGINAS
    content_sources = []
    for i, result in enumerate(results[:4], 1):  # Analisa até 4 páginas
        url = result.get('link', result.get('url', ''))
        title = result.get('title', f'Resultado {i}')
        snippet = result.get('snippet', '')
        
        if url:
            print(f"📄 Extraindo conteúdo de: {title}")
            try:
                # Extrai MUITO mais conteúdo da página
                page_content = extract_page_content(url, max_chars=2000)
                if page_content.strip():
                    content_sources.append({
                        'title': title,
                        'url': url,
                        'content': page_content,
                        'snippet': snippet
                    })
                    print(f"✅ Conteúdo extraído ({len(page_content)} chars)")
                else:
                    print(f"❌ Não foi possível extrair conteúdo de {url}")
            except Exception as e:
                print(f"❌ Erro ao extrair de {url}: {e}")
    
    # SE CONSEGUIU EXTRAIR CONTEÚDO, USA IA PARA ANÁLISE INTELIGENTE
    if content_sources:
        try:
            client = get_text_client()
            if client:
                print("🤖 Processando conteúdo com IA...")
                
                # Monta contexto rico com todo o conteúdo extraído
                context = f"PERGUNTA DO USUÁRIO: {original_query}\n\n"
                context += "CONTEÚDO ENCONTRADO NA INTERNET:\n\n"
                
                for i, source in enumerate(content_sources, 1):
                    context += f"FONTE {i} - {source['title']}\n"
                    context += f"URL: {source['url']}\n"
                    context += f"CONTEÚDO: {source['content']}\n"
                    context += "="*80 + "\n\n"
                
                # Prompt otimizado para resposta direta como ChatGPT
                prompt = f"""Você é a A.E.M.I, especialista em manutenção industrial. Analise o conteúdo extraído da internet e responda DIRETAMENTE à pergunta do usuário.

{context}

INSTRUÇÕES IMPORTANTES:
1. Responda de forma DIRETA e COMPLETA à pergunta
2. Use APENAS as informações encontradas nas fontes
3. Seja técnica e precisa
4. Organize a informação de forma clara
5. NÃO liste as fontes no texto (isso será feito separadamente)
6. Foque em manutenção industrial se aplicável
7. Se a informação não for suficiente, diga isso claramente

RESPOSTA DIRETA:"""

                llm_response = client.text_generation(
                    prompt,
                    max_new_tokens=1200,
                    temperature=0.3,  # Mais conservador para ser mais preciso
                    return_full_text=False
                )
                
                if llm_response and llm_response.strip():
                    print("✅ Resposta IA gerada com sucesso")
                    
                    # Formata resposta final estilo ChatGPT/Gemini
                    final_response = f"🌐 **Resposta baseada em pesquisa na internet:**\n\n"
                    final_response += f"{llm_response.strip()}\n\n"
                    
                    # Adiciona fontes consultadas
                    final_response += "📚 **Fontes consultadas:**\n"
                    for i, source in enumerate(content_sources, 1):
                        final_response += f"{i}. {source['title']}\n"
                        final_response += f"   🔗 {source['url']}\n"
                    
                    return final_response
                else:
                    print("❌ IA não conseguiu gerar resposta")
        
        except Exception as e:
            print(f"❌ Erro ao processar com IA: {e}")
    
    # FALLBACK: Se não conseguiu usar IA, monta resposta básica
    print("⚠️ Usando fallback - resposta básica")
    response = f"🔍 **Pesquisa na Internet:**\n\n"
    
    if content_sources:
        response += f"Encontrei informações sobre **{original_query}**:\n\n"
        
        for i, source in enumerate(content_sources, 1):
            response += f"**{i}. {source['title']}**\n"
            
            # Usa o conteúdo extraído ou snippet
            content_preview = source['content'][:400] if source['content'] else source['snippet']
            if content_preview:
                response += f"📄 {content_preview}...\n"
            
            response += f"🔗 {source['url']}\n\n"
    else:
        # Se não conseguiu extrair conteúdo, usa snippets do Google
        response += f"📊 Encontrei {len(results)} resultado(s) para **{original_query}**:\n\n"
        
        for i, result in enumerate(results, 1):
            title = result.get('title', f'Resultado {i}')
            snippet = result.get('snippet', '')
            url = result.get('link', result.get('url', ''))
            
            response += f"**{i}. {title}**\n"
            if snippet:
                response += f"📋 {snippet}\n"
            response += f"🔗 {url}\n\n"
    
    return response

# --- FUNÇÕES DE PROCESSAMENTO ---
def get_text_client():
    """Cria e retorna um cliente para o modelo de linguagem."""
    if not HUGGING_FACE_TOKEN:
        print("❌ Token HuggingFace não configurado")
        return None
    if not InferenceClient:
        print("❌ InferenceClient não disponível")
        return None
    try:
        client = InferenceClient(model="meta-llama/Meta-Llama-3-8B-Instruct", token=HUGGING_FACE_TOKEN)
        print("✅ Cliente de texto criado com sucesso")
        return client
    except Exception as e:
        print(f"❌ Erro ao criar cliente de texto: {e}")
        return None

def get_vision_client():
    """Cria e retorna um cliente para análise de imagens."""
    if not HUGGING_FACE_TOKEN:
        print("❌ Token HuggingFace não configurado")
        return None
    if not InferenceClient:
        print("❌ InferenceClient não disponível")
        return None
    try:
        client = InferenceClient(model="microsoft/kosmos-2-patch14-224", token=HUGGING_FACE_TOKEN)
        print("✅ Cliente de visão criado com sucesso")
        return client
    except Exception as e:
        print(f"❌ Erro ao criar cliente de visão: {e}")
        return None

def analyze_image(image_path):
    """Analisa uma imagem com IA para uso prático no dia a dia da manutenção."""
    try:
        # 1. ANÁLISE VISUAL COM IA APRIMORADA
        vision_analysis = ""
        practical_analysis = ""
        
        try:
            client = get_vision_client()
            if client and HUGGING_FACE_TOKEN:
                # Converte imagem para base64
                with open(image_path, 'rb') as f:
                    image_data = f.read()
                
                # Prepara a imagem para análise
                img = Image.open(io.BytesIO(image_data))
                
                # Redimensiona se muito grande
                if img.width > 1024 or img.height > 1024:
                    img.thumbnail((1024, 1024), Image.Resampling.LANCZOS)
                
                # Converte para RGB se necessário
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Salva temporariamente
                temp_buffer = io.BytesIO()
                img.save(temp_buffer, format='JPEG', quality=85)
                temp_buffer.seek(0)
                
                # Faz a análise visual com IA
                try:
                    result = client.image_to_text(temp_buffer.getvalue())
                    vision_analysis = result.get('generated_text', '') if isinstance(result, dict) else str(result)
                    
                    # ANÁLISE PRÁTICA COM IA DE TEXTO
                    if vision_analysis and vision_analysis.strip():
                        text_client = get_text_client()
                        if text_client:
                            practical_prompt = f"""Como A.E.M.I, especialista em manutenção industrial, analise esta descrição visual e forneça orientações PRÁTICAS para o dia a dia:

DESCRIÇÃO VISUAL: {vision_analysis}

FORNEÇA UMA ANÁLISE ESTRUTURADA COM:

🔧 **IDENTIFICAÇÃO DO EQUIPAMENTO/COMPONENTE:**
- Que tipo de equipamento/peça identifica?
- Qual a função provável?

⚠️ **CONDIÇÕES OBSERVADAS:**
- Estado atual (bom, desgaste, falha, etc.)
- Sinais de problemas visíveis
- Pontos de atenção

🛠️ **AÇÕES RECOMENDADAS:**
- O que fazer imediatamente
- Procedimentos de segurança
- Quando chamar especialista

💡 **ORIENTAÇÕES PRÁTICAS:**
- Frequência de inspeção
- Sinais para monitorar
- Prevenção de problemas

📋 **PRÓXIMOS PASSOS:**
- Documentação necessária
- Peças/ferramentas necessárias
- Cronograma sugerido

Seja ESPECÍFICA e PRÁTICA para ajudar no dia a dia da manutenção."""

                            try:
                                practical_response = text_client.text_generation(
                                    practical_prompt,
                                    max_new_tokens=1000,
                                    temperature=0.4,
                                    return_full_text=False
                                )
                                if practical_response and practical_response.strip():
                                    practical_analysis = practical_response.strip()
                            except Exception as e:
                                print(f"Erro na análise prática: {e}")
                                
                except Exception as e:
                    print(f"Erro na análise visual: {e}")
                    vision_analysis = ""
        except Exception as e:
            print(f"Erro no cliente de visão: {e}")
            vision_analysis = ""
        
        # 2. ANÁLISE TÉCNICA DA IMAGEM
        with open(image_path, 'rb') as f:
            image_data = f.read()
        
        img = Image.open(io.BytesIO(image_data))
        width, height = img.size
        format_img = img.format or "Desconhecido"
        
        # 3. ANÁLISE BASEADA EM CARACTERÍSTICAS VISUAIS
        visual_characteristics = analyze_visual_characteristics(img)
        
        # 4. ANÁLISE DO NOME DO ARQUIVO
        filename = os.path.basename(image_path).lower()
        context_hints = analyze_filename_context(filename)
        
        # 5. MONTA A RESPOSTA COMPLETA PARA USO PRÁTICO
        if practical_analysis:
            # Resposta com análise prática completa
            description = f"""📸 **ANÁLISE COMPLETA DA IMAGEM:**

{practical_analysis}

📊 **Detalhes técnicos:**
- Formato: {format_img} | Dimensões: {width}x{height}px
{visual_characteristics}
{context_hints}

💬 **Dúvidas?** Pode me perguntar qualquer coisa sobre esta imagem - procedimentos, peças, normas, etc."""
            
        elif vision_analysis:
            # Se temos análise de IA básica, usamos ela
            description = f"""🔍 **Análise Visual da Imagem:**

🤖 **O que identifiquei:**
{vision_analysis}

🔧 **Análise AEMI para seu dia a dia:**
{interpret_for_maintenance(vision_analysis)}

📊 **Características técnicas:**
- Formato: {format_img} | Dimensões: {width}x{height}px
{visual_characteristics}
{context_hints}

💡 **Como posso te ajudar:**
Baseado no que vejo, posso te orientar sobre:
• Identificação de componentes
• Análise de falhas ou desgastes
• Procedimentos de manutenção
• Normas de segurança
• Ferramentas recomendadas

❓ **Próximo passo:** Me conte qual é sua dúvida específica sobre esta imagem."""
        else:
            # Fallback para análise baseada em características
            description = f"""📸 **Análise da Imagem:**

⚠️ **Análise Visual Limitada:**
Não foi possível fazer análise visual completa com IA no momento.

🔧 **Análise AEMI baseada em características:**
{visual_characteristics}
{context_hints}

📊 **Informações técnicas:**
- Formato: {format_img}
- Dimensões: {width}x{height} pixels

💡 **Como posso ajudar:**
Mesmo sem análise visual completa, posso te orientar sobre manutenção industrial se você me descrever:
• Que equipamento/componente está na imagem
• Qual problema você está enfrentando
• Que tipo de análise precisa

❓ **Me conte:** O que você vê na imagem e como posso te ajudar?"""
        
        return description
        
    except Exception as e:
        return f"❌ Erro ao analisar imagem: {str(e)}\n\nTente enviar novamente ou descreva o que você vê na imagem para que eu possa te ajudar."

def analyze_visual_characteristics(img):
    """Analisa características visuais básicas da imagem."""
    try:
        # Análise de cores
        colors = img.getcolors(maxcolors=256*256*256)
        if colors:
            dominant_colors = sorted(colors, key=lambda x: x[0], reverse=True)[:3]
            
            # Interpretação das cores para contexto industrial
            color_hints = []
            for count, color in dominant_colors:
                if isinstance(color, tuple) and len(color) >= 3:
                    r, g, b = color[:3]
                    if r > 200 and g < 100 and b < 100:  # Vermelho
                        color_hints.append("Possível indicação de perigo/parada")
                    elif r > 200 and g > 200 and b < 100:  # Amarelo
                        color_hints.append("Possível sinalização de atenção")
                    elif r < 100 and g > 150 and b < 100:  # Verde
                        color_hints.append("Possível indicação de funcionamento normal")
                    elif r < 100 and g < 100 and b > 150:  # Azul
                        color_hints.append("Possível componente hidráulico")
            
            if color_hints:
                return f"\n🎨 **Indicações visuais:** {', '.join(color_hints)}"
        
        return "\n🎨 **Análise de cores:** Variadas (equipamento/ambiente industrial)"
    except:
        return ""

def analyze_filename_context(filename):
    """Analisa o nome do arquivo para contexto."""
    maintenance_keywords = {
        'motor': 'Motor elétrico/mecânico',
        'rolamento': 'Rolamento/bearing',
        'bearing': 'Rolamento',
        'engrenagem': 'Sistema de engrenagens',
        'gear': 'Engrenagem',
        'bomba': 'Bomba hidráulica/pneumática',
        'pump': 'Bomba',
        'valvula': 'Válvula',
        'valve': 'Válvula',
        'correia': 'Correia/belt',
        'belt': 'Correia',
        'polia': 'Polia',
        'pulley': 'Polia',
        'falha': 'Análise de falha',
        'failure': 'Falha',
        'desgaste': 'Desgaste',
        'wear': 'Desgaste',
        'manutencao': 'Manutenção',
        'maintenance': 'Manutenção',
        'hidraulica': 'Sistema hidráulico',
        'hydraulic': 'Hidráulico',
        'pneumatic': 'Pneumático',
        'pneumatica': 'Pneumático'
    }
    
    found = []
    for keyword, description in maintenance_keywords.items():
        if keyword in filename:
            found.append(description)
    
    if found:
        return f"\n🏷️ **Contexto do arquivo:** {', '.join(found)}"
    return ""

def interpret_for_maintenance(vision_text):
    """Interpreta a análise visual no contexto de manutenção industrial."""
    vision_lower = vision_text.lower()
    
    interpretations = []
    
    # Identifica equipamentos
    if any(word in vision_lower for word in ['motor', 'engine', 'máquina', 'machine']):
        interpretations.append("🔧 **Motor/Máquina identificado** - Posso ajudar com análise de vibração, alinhamento, lubrificação")
    
    if any(word in vision_lower for word in ['rolamento', 'bearing', 'roda', 'wheel']):
        interpretations.append("⚙️ **Rolamento detectado** - Posso orientar sobre montagem, desmontagem e análise de falhas")
    
    if any(word in vision_lower for word in ['tubo', 'pipe', 'mangueira', 'hose']):
        interpretations.append("🔧 **Sistema hidráulico/pneumático** - Posso ajudar com pressões, vedações e conexões")
    
    if any(word in vision_lower for word in ['parafuso', 'bolt', 'rosca', 'thread']):
        interpretations.append("🔩 **Fixação detectada** - Posso orientar sobre torques e procedimentos de aperto")
    
    if any(word in vision_lower for word in ['óleo', 'oil', 'graxa', 'grease', 'lubrificante']):
        interpretations.append("🛢️ **Lubrificação identificada** - Posso ajudar com intervalos e tipos de lubrificantes")
    
    # Identifica problemas
    if any(word in vision_lower for word in ['rachadura', 'crack', 'quebrado', 'broken']):
        interpretations.append("⚠️ **Possível falha estrutural** - Recomendo inspeção detalhada e avaliação de segurança")
    
    if any(word in vision_lower for word in ['oxidação', 'rust', 'corrosão', 'corrosion']):
        interpretations.append("🔴 **Corrosão detectada** - Posso orientar sobre tratamento e prevenção")
    
    if any(word in vision_lower for word in ['desgaste', 'wear', 'gasto', 'worn']):
        interpretations.append("📉 **Desgaste identificado** - Posso ajudar a avaliar vida útil restante")
    
    if interpretations:
        return '\n'.join(interpretations)
    else:
        return "📋 **Análise geral:** Identifiquei elementos industriais. Me descreva sua dúvida específica para orientação detalhada."

def generate_chat_response(chat_history):
    """Processa um histórico de chat e retorna a resposta do modelo."""
    client = get_text_client()
    if not client:
        return "Desculpe, o serviço de IA não está disponível no momento."
    
    try:
        response_generator = client.chat_completion(
            messages=chat_history,
            max_tokens=1500,
            stream=False
        )
        return response_generator.choices[0].message.content
    except Exception as e:
        print(f"Erro na geração de resposta: {e}")
        return "Desculpe, ocorreu um erro ao gerar a resposta."

# --- ROTAS PRINCIPAIS ---
@app.route('/')
def index():
    status = {
        "status": "online",
        "message": "Servidor da AEMI (versão com Llama 3 8B e memória) está no ar.",
        "services": {
            "huggingface": "✅ Configurado" if HUGGING_FACE_TOKEN else "❌ Não configurado",
            "google_search": "✅ Configurado" if (GOOGLE_API_KEY and GOOGLE_CSE_ID) else "❌ Não configurado"
        }
    }
    return jsonify(status)

@app.route('/health')
def health():
    return jsonify({
        "status": "healthy",
        "timestamp": str(__import__('datetime').datetime.now()),
        "google_api": bool(GOOGLE_API_KEY and GOOGLE_CSE_ID),
        "hf_token": bool(HUGGING_FACE_TOKEN)
    })

@app.route('/chat', methods=['POST'])
def chat():
    try:
        user_message = request.form.get('message', '')
        if not user_message.strip():
            return jsonify({"error": "Nenhuma mensagem enviada."}), 400

        # 1. Buscar resposta na base de conhecimento
        kb = load_kb()
        user_message_lower = user_message.lower()
        import difflib
        
        # Busca exata ou substring em FAQ
        for item in kb:
            if item['type'] == 'faq' and user_message_lower in item.get('content','').lower():
                return jsonify({'response': f"[Base de Conhecimento]\n{item['content'][:600]}"})
        
        # Busca fuzzy em FAQ
        best_match = None
        best_ratio = 0.0
        for item in kb:
            if item['type'] == 'faq':
                content = item.get('content','').lower()
                ratio = difflib.SequenceMatcher(None, user_message_lower, content).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_match = item
        
        if best_match and best_ratio > 0.45:
            return jsonify({'response': f"[Base de Conhecimento - resposta aproximada]\n{best_match['content'][:600]}"})

        # 2. Buscar por conteúdo extraído de arquivos
        for item in kb:
            if item['type'] == 'file' and item.get('content') and user_message_lower in item['content'].lower():
                return jsonify({'response': f"[Arquivo: {item['name']}]\n{item['content'][:600]}..."})
        
        # Busca fuzzy no conteúdo de arquivos
        best_file_match = None
        best_file_ratio = 0.0
        for item in kb:
            if item['type'] == 'file' and item.get('content'):
                content = item['content'].lower()
                ratio = difflib.SequenceMatcher(None, user_message_lower, content).ratio()
                if ratio > best_file_ratio:
                    best_file_ratio = ratio
                    best_file_match = item
        
        if best_file_match and best_file_ratio > 0.45:
            return jsonify({'response': f"[Arquivo (aprox.): {best_file_match['name']}]\n{best_file_match['content'][:600]}..."})

        # 3. Buscar por nome de arquivo/documento
        for item in kb:
            if item['type'] == 'file' and user_message_lower in item.get('name','').lower():
                return jsonify({'response': f"Encontrei um documento relacionado: {item['name']}\nClique para baixar: /kb/download/{item['id']}"})
        
        # Fuzzy para nome de arquivo
        best_file = None
        best_file_ratio = 0.0
        for item in kb:
            if item['type'] == 'file':
                name = item.get('name','').lower()
                ratio = difflib.SequenceMatcher(None, user_message_lower, name).ratio()
                if ratio > best_file_ratio:
                    best_file_ratio = ratio
                    best_file = item
        
        if best_file and best_file_ratio > 0.45:
            return jsonify({'response': f"Encontrei um documento relacionado: {best_file['name']}\nClique para baixar: /kb/download/{best_file['id']}"})

        # 4. Verificar se precisa de pesquisa na internet
        if should_search_internet(user_message):
            print(f"Realizando pesquisa na internet para: {user_message}")
            search_results = search_internet(user_message, max_results=5)
            analyzed_results = analyze_search_content(search_results, user_message)
            
            # Se encontrou e analisou resultados, retorna eles
            if "results" in search_results and search_results["results"]:
                return jsonify({"response": analyzed_results})
            
            # Se não encontrou, continua para o LLM com uma nota sobre a pesquisa
            user_message += " (Pesquisa na internet não retornou resultados úteis)"
        
        # 5. Se não encontrou na KB nem precisou pesquisar, usar o LLM
        if 'chat_history' not in session:
            session['chat_history'] = [{"role": "system", "content": SYSTEM_PROMPT}]
        
        # Verifica se há conteúdo de arquivo enviado para incluir na análise
        enhanced_message = user_message
        if 'uploaded_file_content' in session and session['uploaded_file_content']:
            file_data = session['uploaded_file_content']
            
            # Cria contexto baseado no tipo de análise
            if file_data.get('analysis_type') == 'visual':
                enhanced_message = f"""Como A.E.M.I, especialista em manutenção industrial, analise esta imagem e responda: {user_message}

📸 **Arquivo de Imagem:** {file_data['filename']}
🔍 **Análise Visual Completa:**
{file_data['content']}

**Instruções:**
- Responda com base na análise visual da imagem
- Foque em aspectos de manutenção industrial
- Seja técnica e detalhada
- Se identificar problemas, sugira soluções"""
            
            elif file_data.get('analysis_type') == 'text':
                enhanced_message = f"""Como A.E.M.I, especialista em manutenção industrial, analise este documento e forneça orientações PRÁTICAS para o dia a dia. 

PERGUNTA: {user_message}

📄 **Documento:** {file_data['filename']} (tipo: {file_data['type']})

**Conteúdo do Documento:**
{file_data['content']}

**FORNEÇA UMA ANÁLISE ESTRUTURADA COM:**

🔧 **RESUMO EXECUTIVO:**
- Principais pontos do documento
- Informações mais importantes para manutenção

⚠️ **PONTOS CRÍTICOS:**
- Procedimentos de segurança
- Especificações técnicas importantes
- Alertas e cuidados

🛠️ **APLICAÇÃO PRÁTICA:**
- Como usar essas informações no dia a dia
- Procedimentos passo a passo
- Ferramentas necessárias

💡 **DICAS PRÁTICAS:**
- Boas práticas mencionadas
- Frequências recomendadas
- Sinais de alerta

📋 **AÇÕES IMEDIATAS:**
- O que fazer agora
- Próximos passos
- Documentação necessária

Seja ESPECÍFICA e PRÁTICA para ajudar no dia a dia da manutenção."""
            
            else:
                enhanced_message = f"""Pergunta sobre o arquivo enviado: {user_message}

Arquivo: {file_data['filename']} (tipo: {file_data['type']})

Conteúdo do arquivo:
{file_data['content']}

Instruções: Analise o conteúdo do arquivo e responda à pergunta do usuário com base nessas informações."""
        
        session['chat_history'].append({"role": "user", "content": enhanced_message})
        
        if len(session['chat_history']) > MAX_HISTORY_LENGTH:
            session['chat_history'] = [session['chat_history'][0]] + session['chat_history'][-MAX_HISTORY_LENGTH:]
        
        print(f"Processando com histórico de {len(session['chat_history'])} mensagens...")
        if not HUGGING_FACE_TOKEN:
            return jsonify({"response": "⚠️ **Serviço de IA temporariamente indisponível**\n\nPara que eu possa responder adequadamente, é necessário configurar o token da Hugging Face no Replit Secrets.\n\n**Como resolver:**\n1. Clique em 'Secrets' no painel lateral\n2. Adicione a variável `HF_TOKEN` com seu token da Hugging Face\n3. Reinicie a aplicação\n\n**Como obter o token:**\n- Acesse huggingface.co\n- Faça login\n- Vá em Settings → Access Tokens\n- Crie um novo token\n\nApós configurar, estarei pronta para te ajudar com qualquer dúvida sobre manutenção industrial! 🔧"})
        
        bot_response = generate_chat_response(session['chat_history'])
        
        session['chat_history'].append({"role": "assistant", "content": bot_response})
        session.modified = True
        
        print("Resposta da IA gerada e histórico atualizado.")
        return jsonify({"response": bot_response})

    except Exception as e:
        print(f"ERRO GERAL na rota /chat: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Ocorreu um erro inesperado no servidor. Detalhes: {str(e)}"}), 500

@app.route('/upload-file', methods=['POST'])
def upload_file():
    """Upload de arquivo para análise no chat."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
        
        # Salva o arquivo temporariamente
        file_id = str(uuid.uuid4())
        ext = os.path.splitext(file.filename)[1].lower()
        filename = f"{file_id}{ext}"
        temp_path = os.path.join(KB_DIR, filename)
        file.save(temp_path)
        
        # Analisa o arquivo baseado no tipo
        response_text = ""
        file_content = ""
        
        # Verifica se é imagem
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            response_text = "📸 **IMAGEM ANALISADA COM IA - PRONTA PARA TE AJUDAR!**"
            
            # Análise visual da imagem
            try:
                visual_analysis = analyze_image(temp_path)
                if visual_analysis:
                    # Salva análise na sessão
                    session['uploaded_file_content'] = {
                        'filename': file.filename,
                        'type': 'image',
                        'content': visual_analysis,
                        'analysis_type': 'visual'
                    }
                    session.modified = True
                    
                    # Retorna a análise completa automaticamente
                    return jsonify({
                        'response': visual_analysis,
                        'filename': file.filename,
                        'file_type': ext
                    })
            except Exception as e:
                print(f"Erro na análise visual: {e}")
        
        # Verifica se é PDF
        elif ext == '.pdf' and PyPDF2:
            response_text = "📄 **DOCUMENTO ANALISADO COM IA - PRONTO PARA TE AJUDAR!**"
            
            try:
                with open(temp_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text_content = ""
                    
                    # Extrai texto de até 15 páginas
                    pages_to_read = min(15, len(reader.pages))
                    for i in range(pages_to_read):
                        try:
                            page_text = reader.pages[i].extract_text() or ''
                            text_content += f"\n--- Página {i+1} ---\n{page_text}"
                        except:
                            continue
                    
                    # Salva na sessão
                    session['uploaded_file_content'] = {
                        'filename': file.filename,
                        'type': 'pdf',
                        'content': text_content[:15000],  # Aumenta limite
                        'analysis_type': 'text'
                    }
                    session.modified = True
                    
                    # Gera análise automática do documento
                    try:
                        client = get_text_client()
                        if client:
                            doc_analysis_prompt = f"""Como A.E.M.I, especialista em manutenção industrial, analise este documento PDF e forneça um resumo PRÁTICO para uso no dia a dia:

DOCUMENTO: {file.filename}

CONTEÚDO:
{text_content[:10000]}

FORNEÇA:
🔧 **TIPO DE DOCUMENTO:** (manual, procedimento, catálogo, etc.)
📋 **RESUMO EXECUTIVO:** (principais pontos em 3-4 linhas)
⚠️ **PONTOS CRÍTICOS:** (segurança, especificações importantes)
🛠️ **APLICAÇÃO PRÁTICA:** (como usar no dia a dia)
💡 **DICAS IMPORTANTES:** (procedimentos, cuidados, frequências)

Seja DIRETA e PRÁTICA."""

                            doc_analysis = client.text_generation(
                                doc_analysis_prompt,
                                max_new_tokens=800,
                                temperature=0.4,
                                return_full_text=False
                            )
                            
                            if doc_analysis and doc_analysis.strip():
                                full_analysis = f"""📄 **ANÁLISE COMPLETA DO DOCUMENTO:**

{doc_analysis.strip()}

💬 **Como usar:** Agora você pode fazer perguntas específicas sobre qualquer parte do documento. Por exemplo:
• "Como fazer a manutenção preventiva?"
• "Quais são as especificações técnicas?"
• "Que procedimentos de segurança devo seguir?"

❓ **O que você gostaria de saber sobre este documento?**"""
                                
                                return jsonify({
                                    'response': full_analysis,
                                    'filename': file.filename,
                                    'file_type': ext
                                })
                    except Exception as e:
                        print(f"Erro na análise do documento: {e}")
            except Exception as e:
                print(f"Erro ao processar PDF: {e}")
        
        # Verifica se é documento Word
        elif ext in ['.docx'] and docx:
            response_text = "📄 **ARQUIVO RECEBIDO, NO QUE POSSO AJUDAR?**"
            
            try:
                doc = docx.Document(temp_path)
                paragraphs = []
                
                # Extrai parágrafos com formatação
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        paragraphs.append(f"Parágrafo {i+1}: {para.text.strip()}")
                
                text_content = '\n'.join(paragraphs[:30])  # Primeiros 30 parágrafos
                
                # Salva na sessão
                session['uploaded_file_content'] = {
                    'filename': file.filename,
                    'type': 'docx',
                    'content': text_content[:15000],  # Aumenta limite
                    'analysis_type': 'text'
                }
                session.modified = True
                
                # Gera análise automática do documento Word
                try:
                    client = get_text_client()
                    if client:
                        doc_analysis_prompt = f"""Como A.E.M.I, especialista em manutenção industrial, analise este documento Word e forneça um resumo PRÁTICO para uso no dia a dia:

DOCUMENTO: {file.filename}

CONTEÚDO:
{text_content[:10000]}

FORNEÇA:
🔧 **TIPO DE DOCUMENTO:** (manual, procedimento, catálogo, etc.)
📋 **RESUMO EXECUTIVO:** (principais pontos em 3-4 linhas)
⚠️ **PONTOS CRÍTICOS:** (segurança, especificações importantes)
🛠️ **APLICAÇÃO PRÁTICA:** (como usar no dia a dia)
💡 **DICAS IMPORTANTES:** (procedimentos, cuidados, frequências)

Seja DIRETA e PRÁTICA."""

                        doc_analysis = client.text_generation(
                            doc_analysis_prompt,
                            max_new_tokens=800,
                            temperature=0.4,
                            return_full_text=False
                        )
                        
                        if doc_analysis and doc_analysis.strip():
                            full_analysis = f"""📄 **ANÁLISE COMPLETA DO DOCUMENTO:**

{doc_analysis.strip()}

💬 **Como usar:** Agora você pode fazer perguntas específicas sobre qualquer parte do documento. Por exemplo:
• "Como fazer a manutenção preventiva?"
• "Quais são as especificações técnicas?"
• "Que procedimentos de segurança devo seguir?"

❓ **O que você gostaria de saber sobre este documento?**"""
                            
                            return jsonify({
                                'response': full_analysis,
                                'filename': file.filename,
                                'file_type': ext
                            })
                except Exception as e:
                    print(f"Erro na análise do documento: {e}")
            except Exception as e:
                print(f"Erro ao processar documento: {e}")
        
        # Arquivo de texto
        elif ext in ['.txt', '.md', '.csv']:
            response_text = "📝 **ARQUIVO RECEBIDO, NO QUE POSSO AJUDAR?**"
            
            try:
                with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read(15000)  # Aumenta limite
                    
                    # Salva na sessão
                    session['uploaded_file_content'] = {
                        'filename': file.filename,
                        'type': ext,
                        'content': content,
                        'analysis_type': 'text'
                    }
                    session.modified = True
                    
                    # Gera análise automática do arquivo de texto
                    try:
                        client = get_text_client()
                        if client:
                            text_analysis_prompt = f"""Como A.E.M.I, especialista em manutenção industrial, analise este arquivo de texto e forneça um resumo PRÁTICO para uso no dia a dia:

ARQUIVO: {file.filename} (tipo: {ext})

CONTEÚDO:
{content[:10000]}

FORNEÇA:
🔧 **TIPO DE ARQUIVO:** (dados, lista, procedimento, etc.)
📋 **RESUMO EXECUTIVO:** (principais pontos em 3-4 linhas)
⚠️ **PONTOS CRÍTICOS:** (informações importantes)
🛠️ **APLICAÇÃO PRÁTICA:** (como usar no dia a dia)
💡 **DICAS IMPORTANTES:** (procedimentos, cuidados, observações)

Seja DIRETA e PRÁTICA."""

                            text_analysis = client.text_generation(
                                text_analysis_prompt,
                                max_new_tokens=800,
                                temperature=0.4,
                                return_full_text=False
                            )
                            
                            if text_analysis and text_analysis.strip():
                                full_analysis = f"""📝 **ANÁLISE COMPLETA DO ARQUIVO:**

{text_analysis.strip()}

💬 **Como usar:** Agora você pode fazer perguntas específicas sobre qualquer parte do arquivo. Por exemplo:
• "Como interpretar esses dados?"
• "Que procedimentos seguir?"
• "Quais são os pontos mais importantes?"

❓ **O que você gostaria de saber sobre este arquivo?**"""
                                
                                return jsonify({
                                    'response': full_analysis,
                                    'filename': file.filename,
                                    'file_type': ext
                                })
                    except Exception as e:
                        print(f"Erro na análise do arquivo: {e}")
            except Exception as e:
                print(f"Erro ao processar arquivo: {e}")
        
        else:
            response_text = "📎 **ARQUIVO RECEBIDO, NO QUE POSSO AJUDAR?**"
        
        # Remove o arquivo temporário
        try:
            os.remove(temp_path)
        except:
            pass
        
        return jsonify({
            'response': response_text,
            'filename': file.filename,
            'file_type': ext
        })
    
    except Exception as e:
        print(f"Erro no upload de arquivo: {e}")
        return jsonify({'error': 'Erro ao processar arquivo'}), 500

@app.route('/search-internet', methods=['POST'])
def search_internet_route():
    """Rota para pesquisa manual na internet usando Google API."""
    try:
        query = request.form.get('query', '').strip()
        if not query:
            return jsonify({"error": "Query de pesquisa não fornecida"}), 400
        
        max_results = int(request.form.get('max_results', 5))
        
        # Usa diretamente a Google API
        search_results = google_search_api(query, max_results)
        
        # Se houver erro, retorna o erro
        if "error" in search_results:
            return jsonify(search_results), 400
        
        return jsonify(search_results)
    
    except Exception as e:
        print(f"Erro na pesquisa manual: {e}")
        return jsonify({"error": "Erro ao realizar pesquisa"}), 500

@app.route('/extract-content', methods=['POST'])
def extract_content_route():
    """Rota para extrair conteúdo de uma URL."""
    try:
        url = request.form.get('url', '').strip()
        if not url:
            return jsonify({"error": "URL não fornecida"}), 400
        
        max_chars = int(request.form.get('max_chars', 1000))
        content = extract_page_content(url, max_chars)
        
        return jsonify({"content": content, "url": url})
    
    except Exception as e:
        print(f"Erro na extração de conteúdo: {e}")
        return jsonify({"error": "Erro ao extrair conteúdo"}), 500

@app.route('/clear-session', methods=['POST'])
def clear_session():
    """Limpa o histórico de chat da sessão do usuário."""
    try:
        if 'chat_history' in session:
            session.pop('chat_history', None)
            print("Sessão do usuário limpa.")
        return jsonify({"status": "success", "message": "Histórico de chat limpo."})
    except Exception as e:
        print(f"Erro ao limpar sessão: {e}")
        return jsonify({"error": "Erro ao limpar sessão"}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
